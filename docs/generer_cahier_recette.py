# -*- coding: utf-8 -*-
"""Cahier de recette du module Achats — instance de recette."""
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os
import sys

# Ecrit a cote de ce script, dans docs/. Passer un chemin en argument pour
# produire ailleurs : python docs/generer_cahier_recette.py /chemin/vers.docx
DST = sys.argv[1] if len(sys.argv) > 1 else os.path.join(
    os.path.dirname(os.path.abspath(__file__)), 'Cahier_Recette_Module_Achats.docx')

NAVY = '0D1F35'
ENCRE = '1A1A2E'
GRIS = '606A78'
AMBRE = 'B45309'

d = docx.Document()
sec = d.sections[0]
sec.left_margin = sec.right_margin = Cm(1.8)
sec.top_margin = Cm(1.8)
sec.bottom_margin = Cm(1.6)

style = d.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = RGBColor.from_string(ENCRE)


def p(texte='', taille=10.5, gras=False, couleur=ENCRE,
      avant=0, apres=4, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    par = d.add_paragraph()
    par.paragraph_format.space_before = Pt(avant)
    par.paragraph_format.space_after = Pt(apres)
    par.paragraph_format.alignment = align
    if texte:
        r = par.add_run(texte)
        r.bold = gras
        r.font.size = Pt(taille)
        r.font.name = 'Calibri'
        r.font.color.rgb = RGBColor.from_string(couleur)
    return par


def titre1(texte):
    par = p(texte, 15, True, NAVY, avant=18, apres=8, align=WD_ALIGN_PARAGRAPH.LEFT)
    bord = OxmlElement('w:pBdr')
    bas = OxmlElement('w:bottom')
    bas.set(qn('w:val'), 'single')
    bas.set(qn('w:sz'), '10')
    bas.set(qn('w:color'), NAVY)
    bord.append(bas)
    par._p.get_or_add_pPr().append(bord)
    return par


def titre2(texte):
    return p(texte, 11.5, True, ENCRE, avant=12, apres=4, align=WD_ALIGN_PARAGRAPH.LEFT)


def ombrer(cell, couleur):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), couleur)
    cell._tc.get_or_add_tcPr().append(sh)


def tableau(entetes, lignes, largeurs):
    t = d.add_table(rows=1, cols=len(entetes))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # sans mise en page fixe, le rendu redistribue les colonnes a sa guise :
    # la colonne « OK / KO » finissait plus large que le libelle du cas
    t.autofit = False
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    t._tbl.tblPr.append(layout)
    for i, e in enumerate(entetes):
        c = t.rows[0].cells[i]
        ombrer(c, NAVY)
        par = c.paragraphs[0]
        par.paragraph_format.space_after = Pt(2)
        r = par.add_run(e)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string('FFFFFF')
    for ligne in lignes:
        cells = t.add_row().cells
        for i, v in enumerate(ligne):
            par = cells[i].paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            r = par.add_run(str(v))
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string(ENCRE)
    for row in t.rows:
        for i, w in enumerate(largeurs):
            row.cells[i].width = Cm(w)
        # une ligne ne doit pas se couper entre deux pages : un cas de recette
        # scinde en deux est illisible pour celui qui le joue
        pr = row._tr.get_or_add_trPr()
        cs = OxmlElement('w:cantSplit')
        pr.append(cs)
    # l'entete se repete en haut de chaque page
    entete_pr = t.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader')
    entete_pr.append(th)

    # w:tblGrid fait autorite au rendu : sans lui, les largeurs posées sur les
    # cellules sont ignorées et les colonnes reviennent à une répartition
    # égale — la colonne « OK / KO » finissait plus large que le libellé du cas.
    grille = t._tbl.find(qn('w:tblGrid'))
    if grille is not None:
        for col in list(grille):
            grille.remove(col)
        for w in largeurs:
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), str(int(w * 567)))   # cm -> twips
            grille.append(gc)
    return t


def cas(entrees):
    """Tableau de scénarios : N°, action, attendu, verdict, observations."""
    return tableau(
        ['N\u00b0', 'Ce qu\u2019on fait', 'Ce qui doit se produire', 'OK / KO', 'Observations'],
        entrees, [1.2, 5.6, 6.2, 1.6, 3.2])


# ══════════════════════════════════════════════════════════════
p('EMU-CI \u00b7 PROJET NSIIV', 9, True, GRIS, apres=2, align=WD_ALIGN_PARAGRAPH.LEFT)
p('Cahier de recette', 24, True, NAVY, apres=2, align=WD_ALIGN_PARAGRAPH.LEFT)
p('Module Achats \u2014 ERP EMUCI', 14, False, GRIS, apres=10, align=WD_ALIGN_PARAGRAPH.LEFT)
p('Version 1 \u2014 21/08/2026 \u2014 instance de recette, branche recette-achats', 9, False, GRIS, apres=14)

p("Ce cahier sert \u00e0 \u00e9prouver le module Achats de bout en bout, sur une instance d\u00e9di\u00e9e. "
  "Chaque cas se joue dans l'ordre : les suivants s'appuient sur ce que les pr\u00e9c\u00e9dents ont produit. "
  "Cochez OK ou KO, et notez ce que vous avez vu \u2014 surtout en cas de KO : ce que vous attendiez, "
  "ce qui s'est affich\u00e9, et l'heure. C'est cette derni\u00e8re qui permet de retrouver la trace.")

titre1("Se connecter")
p("L'instance est \u00e0 l'adresse suivante. Elle est distincte de la production : "
  "rien de ce que vous y ferez n'a d'effet sur les donn\u00e9es r\u00e9elles.")
p("https://erp-emuci-recette.onrender.com", 12, True, AMBRE, apres=8, align=WD_ALIGN_PARAGRAPH.LEFT)

p("Tous les comptes partagent le m\u00eame mot de passe :", apres=2)
p("Recette@2026", 12, True, AMBRE, apres=8, align=WD_ALIGN_PARAGRAPH.LEFT)

tableau(['Compte', 'R\u00f4le', 'Ce qu\u2019il fait dans le parcours'],
        [
         ('testoperation@gmail.com', 'Demandeur (Op\u00e9rations)', "Cr\u00e9e les demandes d'achat."),
         ('resp.operations@recette.local', 'Responsable N+1 Op\u00e9rations', "Endosse les demandes de son service."),
         ('achat@recette.local', 'Acheteur', "Prend en charge, arbitre, saisit les offres, lance la validation."),
         ('raf@recette.local', 'RAF', "Premier signataire financier. Endosse aussi pour l'Administration."),
         ('daf@recette.local', 'DAF', "Deuxi\u00e8me signataire, au-del\u00e0 de 500 000 XOF."),
         ('pdg@recette.local', 'Direction g\u00e9n\u00e9rale', "Troisi\u00e8me signataire, au-del\u00e0 de 5 000 000 XOF."),
         ('magasin@recette.local', 'Magasin', "Exp\u00e9die les commandes internes, enregistre les r\u00e9ceptions."),
        ], [6.0, 4.2, 7.6])

titre2("Trois choses \u00e0 savoir avant de commencer")
p("\u2022 Le premier \u00e9cran met environ cinquante secondes \u00e0 s'ouvrir quand l'instance dort. "
  "Ce n'est pas une panne : elle s'endort apr\u00e8s un quart d'heure sans visite.", apres=2)
p("\u2022 Vous \u00eates d\u00e9connect\u00e9 au bout de quinze minutes sans action. C'est voulu, et c'est "
  "l'objet du cas T-02.", apres=2)
p("\u2022 Les fichiers que vous d\u00e9posez disparaissent \u00e0 chaque red\u00e9ploiement de l'instance. "
  "Si un document t\u00e9l\u00e9vers\u00e9 la veille a disparu, ce n'est pas une anomalie \u2014 c'est une limite "
  "connue de l'h\u00e9bergement d'essai.", apres=8)

titre1("Le jeu de donn\u00e9es fourni")
tableau(['\u00c9l\u00e9ment', 'Ce qui est en place'],
        [
         ('Stock du magasin central', 'rivets 500, papier toilette 80, caf\u00e9 40, sucre 40, LIPTON 30, '
                                      'Lotus 30, huile 25'),
         ('Fournisseurs', 'DMD et INFOSOLUCES \u2014 volontairement sans pi\u00e8ces de conformit\u00e9'),
         ('Paliers de signature', 'jusqu\u2019\u00e0 500 000 : RAF \u00b7 jusqu\u2019\u00e0 5 000 000 : RAF + DAF \u00b7 au-del\u00e0 : '
                                  'RAF + DAF + PDG'),
         ('Budget', 'Six lignes pour 2026, toutes sans enveloppe, en mode alerte, et rattachées à aucun service. Aucun contrôle n’est donc actif : la famille I commence par en créer un.'),
         ('Nomenclatures d\u2019\u00e9quipement', 'clavier, imprimante, serveur, unit\u00e9 centrale, \u00e9cran, onduleur, '
                                            'portable, presse'),
        ], [4.6, 13.2])

# ══════════════════════════════════════════════════════════════
titre1("A \u00b7 Expression du besoin")
p("Compte : testoperation@gmail.com", 9.5, True, GRIS, apres=6, align=WD_ALIGN_PARAGRAPH.LEFT)
cas([
 ('A-01', "Ouvrir \u00ab Nouvelle FEB \u00bb.",
  "Le site et le service sont d\u00e9j\u00e0 remplis d'apr\u00e8s le profil.", '', ''),
 ('A-02', "Soumettre sans aucune ligne.",
  "Refus : au moins une ligne est exig\u00e9e.", '', ''),
 ('A-03', "Ajouter une ligne \u00ab rivets \u00bb, quantit\u00e9 50, sans choisir de famille, puis soumettre.",
  "Refus : la famille est obligatoire, la ligne concern\u00e9e est nomm\u00e9e.", '', ''),
 ('A-04', "Compl\u00e9ter la famille et le type d'achat, puis enregistrer en brouillon.",
  "Brouillon enregistr\u00e9, aucun num\u00e9ro attribu\u00e9.", '', ''),
 ('A-05', "Ajouter trois lignes de familles diff\u00e9rentes, puis une quatri\u00e8me d'une famille nouvelle.",
  "La quatri\u00e8me est refus\u00e9e : trois familles au maximum par demande.", '', ''),
 ('A-06', "Soumettre la demande (garder trois familles au plus).",
  "Num\u00e9ro FEB-2026-NNNN attribu\u00e9. La demande appara\u00eet dans \u00ab Mes FEB \u00bb.", '', ''),
])

titre1("B \u00b7 Prise en charge et arbitrage sur stock")
p("Compte : achat@recette.local", 9.5, True, GRIS, apres=6, align=WD_ALIGN_PARAGRAPH.LEFT)
cas([
 ('B-01', "Ouvrir \u00ab File d'attente Achats \u00bb.",
  "La demande de A-06 y figure, avec son urgence et son anciennet\u00e9. Aucun message d'erreur "
  "en haut de page.", '', ''),
 ('B-02', "Cliquer \u00ab Prendre en charge \u00bb.",
  "Votre nom s'affiche comme preneur. La demande quitte la liste \u00ab \u00e0 prendre en charge \u00bb.", '', ''),
 ('B-03', "Se connecter avec un autre compte acheteur et tenter de la reprendre.",
  "Refus : elle est d\u00e9j\u00e0 prise en charge.", '', ''),
 ('B-04', "Sur la ligne \u00ab rivets \u00bb, choisir \u00ab Servir sur stock \u00bb.",
  "Accept\u00e9 : le magasin en d\u00e9tient 500. Un avertissement signale que le stock n'est pas sur "
  "le site demandeur et suppose un transfert.", '', ''),
 ('B-05', "Basculer vers la commande interne.",
  "Une commande CMD-... est cr\u00e9\u00e9e, avec la mention \u00ab Issue de la FEB ... \u00bb. "
  "Le stock du magasin n'a pas encore boug\u00e9.", '', ''),
])

titre1("C \u00b7 Offres et conformit\u00e9 des fournisseurs")
p("Compte : achat@recette.local", 9.5, True, GRIS, apres=6, align=WD_ALIGN_PARAGRAPH.LEFT)
cas([
 ('C-01', "Sur une ligne rest\u00e9e en achat, saisir une offre DMD \u00e0 300 000 XOF.",
  "Offre enregistr\u00e9e et affich\u00e9e dans le comparatif du lot.", '', ''),
 ('C-02', "Cliquer deux fois de suite sur \u00ab Enregistrer \u00bb dans la fen\u00eatre d'offre.",
  "Une seule offre est cr\u00e9\u00e9e. Le bouton se d\u00e9sactive pendant l'envoi.", '', ''),
 ('C-03', "Saisir une deuxi\u00e8me offre INFOSOLUCES, puis une troisi\u00e8me, puis tenter une quatri\u00e8me.",
  "La quatri\u00e8me est refus\u00e9e : trois offres au maximum par lot.", '', ''),
 ('C-04', "Cliquer \u00ab Retenir \u00bb sur l'offre DMD.",
  "Refus : le dossier de conformit\u00e9 de DMD est incomplet. Le message nomme les pi\u00e8ces "
  "manquantes \u2014 RCCM, DFE, RIB, PIRL.", '', ''),
 ('C-05', "Aller dans Achats \u2192 Fournisseurs, ouvrir DMD, d\u00e9poser les quatre pi\u00e8ces obligatoires "
          "(n'importe quel PDF fait l'affaire).",
  "Les quatre pi\u00e8ces s'affichent. Le bouton \u00ab Documents \u00bb de la liste permet de les consulter "
  "sans quitter l'\u00e9cran.", '', ''),
 ('C-06', "Revenir sur la demande et retenir l'offre DMD.",
  "Accept\u00e9 : le fournisseur est report\u00e9 sur les lignes du lot.", '', ''),
 ('C-07', "Sur une ligne, choisir \u00e0 la main le fournisseur INFOSOLUCES.",
  "Refus \u00e9galement : la r\u00e8gle vaut pour les deux fa\u00e7ons de d\u00e9signer un fournisseur.", '', ''),
 ('C-08', "Saisir le montant de la ligne, puis \u00ab V\u00e9rifier le comparatif \u00bb.",
  "Le contr\u00f4le passe si la somme des lignes \u00e9gale le montant de l'offre retenue ; sinon il "
  "affiche les deux chiffres.", '', ''),
])

titre1("D \u00b7 Endossement et visas")
p("Comptes : achat, resp.operations, raf, daf, pdg", 9.5, True, GRIS, apres=6, align=WD_ALIGN_PARAGRAPH.LEFT)
cas([
 ('D-01', "Acheteur : lancer la validation sur une demande \u00e0 300 000 XOF.",
  "Le circuit annonc\u00e9 est \u00ab Responsable N+1, puis RAF \u00bb \u2014 une seule signature financi\u00e8re "
  "sous 500 000.", '', ''),
 ('D-02', "RAF : ouvrir \u00ab Mes visas \u00bb.",
  "La demande n'y est pas : l'endossement du responsable la pr\u00e9c\u00e8de.", '', ''),
 ('D-03', "resp.operations : ouvrir \u00ab Mes visas \u00bb.",
  "La demande y figure. L'\u00e9cran s'ouvre bien, alors que ce compte n'est pas un valideur "
  "financier.", '', ''),
 ('D-04', "resp.operations : examiner la demande.",
  "Le d\u00e9tail montre le budget du service par famille, le lot, et les offres avec la retenue "
  "signal\u00e9e.", '', ''),
 ('D-05', "resp.operations : accepter, avec un commentaire.",
  "Visa enregistr\u00e9. La demande passe au RAF, qui est notifi\u00e9.", '', ''),
 ('D-06', "RAF : accepter.",
  "La demande est confirm\u00e9e automatiquement \u2014 il n'y a pas d'acte de confirmation s\u00e9par\u00e9 "
  "\u00e0 accomplir.", '', ''),
 ('D-07', "Refaire une demande \u00e0 2 000 000 XOF et la mener jusqu'aux visas.",
  "Le circuit compte cette fois le RAF puis le DAF. Le DAF ne voit rien tant que le RAF "
  "n'a pas sign\u00e9.", '', ''),
 ('D-08', "Refaire \u00e0 9 000 000 XOF.",
  "Le circuit ajoute la Direction g\u00e9n\u00e9rale en troisi\u00e8me signature.", '', ''),
 ('D-09', "Sur une de ces demandes, faire refuser un signataire sans motif.",
  "Refus impossible : le motif est obligatoire.", '', ''),
 ('D-10', "Refuser avec motif.",
  "La demande revient \u00e0 l'acheteur, le motif est conserv\u00e9 et visible.", '', ''),
 ('D-11', "Demandeur : tenter de viser sa propre demande.",
  "Impossible : nul ne vise sa propre demande.", '', ''),
])

titre1("E \u00b7 R\u00e9f\u00e9rences Sage")
p("Compte : achat@recette.local", 9.5, True, GRIS, apres=6, align=WD_ALIGN_PARAGRAPH.LEFT)
cas([
 ('E-01', "Ouvrir \u00ab Suivi Achats \u00bb apr\u00e8s une confirmation.",
  "Une ligne de suivi existe pour chaque ligne achet\u00e9e \u2014 aucune pour les lignes servies "
  "sur stock.", '', ''),
 ('E-02', "Tenter de saisir un N\u00b0 BC avant le N\u00b0 DA.",
  "Le champ BC est indisponible et signale que la DA est requise.", '', ''),
 ('E-03', "Saisir un N\u00b0 DA, puis un N\u00b0 BC et une date de livraison pr\u00e9vue.",
  "Les deux r\u00e9f\u00e9rences sont enregistr\u00e9es, le statut de la ligne se recalcule seul.", '', ''),
 ('E-04', "Utiliser \u00ab Appliquer au lot \u00bb sur une demande \u00e0 plusieurs lignes.",
  "La r\u00e9f\u00e9rence est pos\u00e9e sur toutes les lignes du lot en une seule action.", '', ''),
])

titre1("F \u00b7 R\u00e9ception")
p("Compte : magasin@recette.local ou achat@recette.local", 9.5, True, GRIS, apres=6, align=WD_ALIGN_PARAGRAPH.LEFT)
cas([
 ('F-01', "Ouvrir \u00ab R\u00e9ceptions \u00bb.",
  "Seules les lignes portant un N\u00b0 BC y figurent.", '', ''),
 ('F-02', "R\u00e9ceptionner une partie de la quantit\u00e9 attendue.",
  "L'\u00e9cart est conserv\u00e9, la ligne reste ouverte, la date de livraison r\u00e9elle n'est pas "
  "encore pos\u00e9e.", '', ''),
 ('F-03', "Tenter de recevoir plus que le reste attendu.",
  "Refus chiffr\u00e9 : le message indique le reste et la quantit\u00e9 saisie.", '', ''),
 ('F-04', "R\u00e9ceptionner le solde.",
  "La ligne est sold\u00e9e, la date r\u00e9elle est pos\u00e9e, le demandeur re\u00e7oit un \u00ab besoin couvert \u00bb.", '', ''),
 ('F-05', "V\u00e9rifier le stock de l'article re\u00e7u.",
  "Le stock global, celui du site de la demande et celui du service ont augment\u00e9 de la "
  "quantit\u00e9 re\u00e7ue.", '', ''),
])

titre1("G \u00b7 Commande interne et stock du magasin")
p("Comptes : superviseur Op\u00e9rations, magasin, puis le site destinataire", 9.5, True, GRIS,
  apres=6, align=WD_ALIGN_PARAGRAPH.LEFT)
cas([
 ('G-01', "Ouvrir la commande interne cr\u00e9\u00e9e en B-05 et la valider ligne par ligne.",
  "La commande passe \u00ab \u00e0 pr\u00e9parer \u00bb.", '', ''),
 ('G-02', "Noter le stock du magasin, puis exp\u00e9dier une quantit\u00e9 inf\u00e9rieure \u00e0 celle demand\u00e9e.",
  "Un motif d'\u00e9cart est exig\u00e9 avant de pouvoir exp\u00e9dier.", '', ''),
 ('G-03', "Exp\u00e9dier, puis revenir sur le stock du magasin.",
  "Le magasin a bien \u00e9t\u00e9 d\u00e9bit\u00e9 de la quantit\u00e9 exp\u00e9di\u00e9e \u2014 c'est le point corrig\u00e9 le 18/08.", '', ''),
 ('G-04', "R\u00e9ceptionner la commande c\u00f4t\u00e9 site destinataire.",
  "Le stock du site destinataire augmente d'autant.", '', ''),
])

titre1("H \u00b7 Immobilisations et affectation")
p("Comptes : achat, magasin, puis les signataires", 9.5, True, GRIS, apres=6, align=WD_ALIGN_PARAGRAPH.LEFT)
cas([
 ('H-01', "Cr\u00e9er une demande d'un onduleur, type d'achat DAI, quantit\u00e9 2, et la mener "
          "jusqu'\u00e0 la confirmation.",
  "Parcours normal, rien de particulier \u00e0 ce stade.", '', ''),
 ('H-02', "Dans \u00ab R\u00e9ceptions \u00bb, rattacher la ligne \u00e0 la nomenclature \u00ab onduleur \u00bb.",
  "Le rattachement est accept\u00e9. Il reste facultatif : une ligne non rattach\u00e9e se re\u00e7oit "
  "normalement.", '', ''),
 ('H-03', "R\u00e9ceptionner les deux unit\u00e9s.",
  "Deux exemplaires sont cr\u00e9\u00e9s \u2014 un par unit\u00e9 \u2014 valoris\u00e9s chacun \u00e0 la moiti\u00e9 du montant "
  "de la ligne.", '', ''),
 ('H-04', "Ouvrir la file d'attente d'affectation.",
  "Les deux exemplaires y figurent, en attente d'affectation, rattach\u00e9s au service qui "
  "les a pay\u00e9s.", '', ''),
 ('H-05', "Proposer l'affectation d'un exemplaire \u00e0 un site.",
  "Un circuit de validation s'ouvre, d\u00e9termin\u00e9 par la valeur de l'exemplaire \u2014 les m\u00eames "
  "paliers que pour les achats.", '', ''),
 ('H-06', "Faire refuser l'affectation.",
  "L'exemplaire revient en attente d'affectation.", '', ''),
 ('H-07', "Reproposer et faire accepter.",
  "L'exemplaire porte d\u00e9sormais son site, et son statut passe \u00e0 affect\u00e9.", '', ''),
])

titre1("I \u00b7 Contr\u00f4le budg\u00e9taire")
p("Comptes : raf, pdg, achat", 9.5, True, GRIS, apres=6, align=WD_ALIGN_PARAGRAPH.LEFT)
ic = [
 ('I-01', "RAF : ouvrir « Lignes budgétaires ». Créer une ligne pour le service Opérations, "
          "famille Fournitures d'entretien, exercice 2026, enveloppe 1 000 000, comportement « blocage ».",
  "La ligne apparaît. Le budget du service est à l'état brouillon.", '', ''),
 ('I-02', "RAF : soumettre le budget du service Opérations.",
  "Le budget passe « soumis » et la direction est notifiée.", '', ''),
 ('I-03', "PDG : valider le budget.",
  "Le budget passe « validé ».", '', ''),
 ('I-04', "RAF : tenter de modifier l'enveloppe qui vient d'être validée.",
  "Refus : un budget validé est verrouillé.", '', ''),
 ('I-05', "Créer une demande d'entretien de 1 500 000 XOF pour les Opérations, puis lancer la validation.",
  "Le lancement est refusé : le dépassement est bloquant sur cette ligne.", '', ''),
 ('I-06', "Refaire la même chose sur une famille sans enveloppe — les consommables informatiques.",
  "Un avertissement s'affiche — « non contrôlé » — mais la validation passe : un budget "
  "non arrêté ne doit jamais paralyser le service Achats.", '', ''),
 ('I-07', "RAF : basculer le filtre sur l'exercice 2027 et y créer une ligne.",
  "L'année à venir est ouverte à la saisie.", '', ''),
 ('I-08', "RAF : utiliser l'import en masse avec un petit fichier CSV de deux lignes.",
  "Un aperçu s'affiche avant écriture. L'import respecte le verrouillage d'un budget déjà validé.", '', ''),
]
cas(ic)

titre1("T \u00b7 Transverses")
cas([
 ('T-01', "Demandeur : tenter d'ouvrir la file d'attente Achats ou le param\u00e9trage.",
  "Acc\u00e8s refus\u00e9 : ces \u00e9crans sont r\u00e9serv\u00e9s au service Achats.", '', ''),
 ('T-02', "Laisser un \u00e9cran ouvert quinze minutes sans y toucher, puis cliquer.",
  "Vous \u00eates ramen\u00e9 \u00e0 la connexion, avec la mention que la session a expir\u00e9.", '', ''),
 ('T-03', "Sur une demande confirm\u00e9e, ouvrir la fiche imprimable.",
  "Elle porte un bloc \u00ab Demandeur / Sup\u00e9rieur hi\u00e9rarchique \u00bb, les num\u00e9ros DA et BC, le "
  "comparatif avec une croix sur l'offre retenue, et le contr\u00f4le budg\u00e9taire.", '', ''),
 ('T-04', "Ouvrir la fiche de validation archiv\u00e9e.",
  "Elle ne porte que les signataires financiers \u2014 le responsable N+1 n'y figure pas.", '', ''),
 ('T-05', "Parcourir les \u00e9crans du module et signaler tout message d'erreur, cadre vide "
          "ou bouton sans effet.",
  "Aucun avertissement PHP, aucun bouton inerte.", '', ''),
])

titre1("Ce qui est connu, et qu'il est inutile de signaler")
tableau(['Constat', 'Pourquoi'],
        [
         ("L'endossement du responsable a lieu au moment o\u00f9 l'acheteur lance la validation, "
          "et non d\u00e8s la soumission.",
          "\u00c9cart connu. Il a \u00e9t\u00e9 d\u00e9cid\u00e9 de d\u00e9placer cet endossement \u00e0 la soumission ; "
          "le d\u00e9veloppement reste \u00e0 faire."),
         ("Les fichiers d\u00e9pos\u00e9s disparaissent apr\u00e8s un red\u00e9ploiement.",
          "Limite de l'h\u00e9bergement d'essai, sans disque persistant."),
         ("Le premier \u00e9cran met pr\u00e8s d'une minute \u00e0 s'ouvrir.",
          "L'instance s'endort apr\u00e8s un quart d'heure sans visite."),
         ("Les sch\u00e9mas de la sp\u00e9cification montrent des \u00e9tapes et des statuts qui ne "
          "correspondent plus.",
          "Ils datent de la premi\u00e8re version et seront repris ; le texte fait foi."),
         ("Deux responsables peuvent \u00eatre d\u00e9sign\u00e9s pour un m\u00eame service.",
          "Rien ne l'interdit aujourd'hui ; un seul sera sollicit\u00e9. \u00c0 trancher."),
        ], [8.0, 9.8])

titre1("Bordereau de synth\u00e8se")
p("\u00c0 remplir en fin de recette, puis \u00e0 renvoyer avec le cahier annot\u00e9.", apres=8)
tableau(['Famille', 'Cas jou\u00e9s', 'OK', 'KO', 'Remarques'],
        [('A \u00b7 Expression du besoin', '6', '', '', ''),
         ('B \u00b7 Prise en charge et arbitrage', '5', '', '', ''),
         ('C \u00b7 Offres et conformit\u00e9', '8', '', '', ''),
         ('D \u00b7 Endossement et visas', '11', '', '', ''),
         ('E \u00b7 R\u00e9f\u00e9rences Sage', '4', '', '', ''),
         ('F \u00b7 R\u00e9ception', '5', '', '', ''),
         ('G \u00b7 Commande interne', '4', '', '', ''),
         ('H \u00b7 Immobilisations', '7', '', '', ''),
         ('I \u00b7 Budget', '5', '', '', ''),
         ('T \u00b7 Transverses', '5', '', '', ''),
         ('Total', '63', '', '', ''),
         ], [5.4, 2.2, 1.6, 1.6, 7.0])

p()
tableau(['Recette prononc\u00e9e par', 'Date', 'Avis'],
        [('', '', 'Accept\u00e9e \u00b7 Accept\u00e9e sous r\u00e9serve \u00b7 Refus\u00e9e'),
         ('', '', '')], [6.0, 3.4, 8.4])

pied = sec.footer.paragraphs[0]
pied.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pied.add_run('EMU-CI \u00b7 Projet NSIIV \u00b7 Cahier de recette Module Achats \u2014 v1.0')
r.font.size = Pt(8)
r.font.color.rgb = RGBColor.from_string(GRIS)

d.save(DST)
print('cahier ecrit :', DST)
